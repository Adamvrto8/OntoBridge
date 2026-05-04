from __future__ import annotations

from pathlib import Path

from ontobridge.agents.harvester.protocols import RawDocument
from ontobridge.models.enums import SourceType

# Heading styles used by Word to mark section titles
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3",
    "heading 4", "heading 5", "heading 6",
    "title", "subtitle",
}


class DocxReader:
    """Reads .docx files using python-docx.

    Preserves heading structure: a heading paragraph becomes the ``section``
    of subsequent body paragraphs.  When a short heading is followed by a
    longer paragraph that reads like a definition, both are emitted together
    so the PatternTermExtractor can treat heading=term, body=definition.

    Requires: pip install python-docx>=1.0  (or the [readers] optional dep group)
    """

    source_type: SourceType = SourceType.POLICY_DOC

    def can_read(self, source: Path | str) -> bool:
        return Path(source).suffix.lower() in {".docx", ".doc"}

    def read(self, source: Path | str) -> list[RawDocument]:
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DocxReader.\n"
                "Install it with:  pip install python-docx>=1.0"
            ) from exc

        document = docx.Document(str(source))
        docs: list[RawDocument] = []
        current_section: str | None = None
        pending_heading: str | None = None  # short heading that may pair with next para

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = (para.style.name or "").lower()
            is_heading = style in _HEADING_STYLES

            if is_heading:
                current_section = text
                pending_heading = text
                continue

            # Short title-case line without a heading style but looks like a term name
            words = text.split()
            looks_like_term = (
                len(words) <= 6
                and text[0].isupper()
                and not text.endswith(".")
                and ":" not in text
            )
            if looks_like_term:
                pending_heading = text
                current_section = text
                continue

            # Body paragraph — emit it, prepending pending heading if present
            if pending_heading:
                # Emit as "TERM: definition text" so PatternTermExtractor can parse it
                combined = f"{pending_heading}: {text}"
                docs.append(RawDocument(text=combined, section=current_section))
                pending_heading = None
            else:
                docs.append(RawDocument(text=text, section=current_section))

        return docs
